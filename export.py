from docx import Document as DocxDocument
from database import SessionLocal
import models


def export_answers_to_docx(filename="output.docx"):

    db = SessionLocal()
    answers = db.query(models.Answer).order_by(models.Answer.id).all()

    doc = DocxDocument()

    for ans in answers:
        doc.add_heading(ans.question, level=2)
        doc.add_paragraph("Answer:")
        doc.add_paragraph(ans.answer)
        doc.add_paragraph(f"Citation: {ans.citation}")
        doc.add_paragraph(" ")

    doc.save(filename)

    return filename